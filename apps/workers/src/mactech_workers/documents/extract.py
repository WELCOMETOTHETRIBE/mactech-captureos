"""Multi-format text extraction.

Ordinary text extraction first; controlled OCR only when a PDF yields no
embedded text (a scanned image). Every extractor is defensive: a parse failure
returns an empty ``ExtractedDoc`` with ``ok=False`` rather than raising, so one
bad attachment never sinks a whole procurement package.
"""

from __future__ import annotations

import csv
import io
import logging
from dataclasses import dataclass, field
from html.parser import HTMLParser

log = logging.getLogger(__name__)

OCR_FALLBACK_THRESHOLD = 80
OCR_RENDER_DPI = 220
OCR_MAX_PAGES = 12
MAX_TEXT_CHARS = 200_000  # per-document cap (generous; whole-package cap is upstream)

# format -> canonical mime
_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv",
    "txt": "text/plain",
    "html": "text/html",
    "zip": "application/zip",
    "unknown": "application/octet-stream",
}


@dataclass
class ExtractedDoc:
    text: str
    format: str
    mime_type: str
    ok: bool
    pages: list[str] = field(default_factory=list)  # per-page for PDF; else [text]
    page_count: int | None = None
    ocr_used: bool = False
    error: str | None = None


def detect_format(filename: str, blob: bytes) -> str:
    name = (filename or "").lower()
    if blob[:4] == b"%PDF" or name.endswith(".pdf"):
        return "pdf"
    if blob[:4] in (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08"):
        # OOXML files are ZIPs; disambiguate by extension.
        if name.endswith(".docx"):
            return "docx"
        if name.endswith((".xlsx", ".xlsm")):
            return "xlsx"
        if name.endswith((".pptx",)):
            return "unknown"
        return "zip"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith((".xlsx", ".xls", ".xlsm")):
        return "xlsx"
    if name.endswith(".csv"):
        return "csv"
    if name.endswith((".htm", ".html")):
        return "html"
    if name.endswith((".txt", ".md")):
        return "txt"
    # Sniff text vs binary.
    sample = blob[:512]
    if sample and _looks_texty(sample):
        if b"<html" in sample.lower() or b"<!doctype html" in sample.lower():
            return "html"
        return "txt"
    return "unknown"


def _looks_texty(sample: bytes) -> bool:
    if b"\x00" in sample:
        return False
    try:
        sample.decode("utf-8")
        return True
    except UnicodeDecodeError:
        try:
            sample.decode("latin-1")
            return True
        except UnicodeDecodeError:
            return False


def _decode(blob: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return blob.decode(enc)
        except UnicodeDecodeError:
            continue
    return blob.decode("utf-8", errors="replace")


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs: object) -> None:
        if tag in ("script", "style"):
            self._skip += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style") and self._skip:
            self._skip -= 1

    def handle_data(self, data: str) -> None:
        if not self._skip and data.strip():
            self._chunks.append(data.strip())

    @property
    def text(self) -> str:
        return "\n".join(self._chunks)


def _extract_pdf(blob: bytes, *, allow_ocr: bool) -> ExtractedDoc:
    import fitz  # type: ignore[import-untyped]

    pages: list[str] = []
    try:
        with fitz.open(stream=blob, filetype="pdf") as doc:
            page_count = doc.page_count
            for page in doc:
                pages.append(page.get_text("text"))
    except Exception as exc:  # fitz.FileDataError, RuntimeError, …
        return ExtractedDoc(text="", format="pdf", mime_type=_MIME["pdf"], ok=False, error=str(exc))

    embedded = "\n\n".join(p.strip() for p in pages).strip()
    if len(embedded) >= OCR_FALLBACK_THRESHOLD or not allow_ocr:
        return ExtractedDoc(
            text=embedded[:MAX_TEXT_CHARS],
            format="pdf",
            mime_type=_MIME["pdf"],
            ok=bool(embedded),
            pages=pages,
            page_count=page_count,
        )

    ocr_pages = _ocr_pdf(blob)
    ocr_text = "\n\n".join(ocr_pages).strip()
    if len(ocr_text) > len(embedded):
        return ExtractedDoc(
            text=ocr_text[:MAX_TEXT_CHARS],
            format="pdf",
            mime_type=_MIME["pdf"],
            ok=bool(ocr_text),
            pages=ocr_pages or pages,
            page_count=page_count,
            ocr_used=True,
        )
    return ExtractedDoc(
        text=embedded[:MAX_TEXT_CHARS],
        format="pdf",
        mime_type=_MIME["pdf"],
        ok=bool(embedded),
        pages=pages,
        page_count=page_count,
    )


def _ocr_pdf(blob: bytes) -> list[str]:
    try:
        import fitz  # type: ignore[import-untyped]
        import pytesseract  # type: ignore[import-untyped]
        from PIL import Image
    except Exception as exc:  # pragma: no cover - optional OCR stack
        log.warning("OCR stack unavailable: %s", exc)
        return []

    pages: list[str] = []
    try:
        with fitz.open(stream=blob, filetype="pdf") as doc:
            for i, page in enumerate(doc):
                if i >= OCR_MAX_PAGES:
                    break
                pix = page.get_pixmap(dpi=OCR_RENDER_DPI, alpha=False)
                with Image.open(io.BytesIO(pix.tobytes("png"))) as img:
                    text = pytesseract.image_to_string(img, lang="eng")
                pages.append(text.strip())
    except Exception as exc:
        log.warning("OCR fall-through errored: %s", exc)
        return []
    return pages


def _extract_docx(blob: bytes) -> ExtractedDoc:
    try:
        import docx  # type: ignore[import-untyped]

        document = docx.Document(io.BytesIO(blob))
        paras = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        text = "\n".join(paras).strip()
        return ExtractedDoc(
            text=text[:MAX_TEXT_CHARS],
            format="docx",
            mime_type=_MIME["docx"],
            ok=bool(text),
            pages=[text],
        )
    except Exception as exc:
        return ExtractedDoc(
            text="", format="docx", mime_type=_MIME["docx"], ok=False, error=str(exc)
        )


def _extract_xlsx(blob: bytes) -> ExtractedDoc:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append("\t".join(cells))
        wb.close()
        text = "\n".join(lines).strip()
        return ExtractedDoc(
            text=text[:MAX_TEXT_CHARS],
            format="xlsx",
            mime_type=_MIME["xlsx"],
            ok=bool(text),
            pages=[text],
        )
    except Exception as exc:
        return ExtractedDoc(
            text="", format="xlsx", mime_type=_MIME["xlsx"], ok=False, error=str(exc)
        )


def _extract_csv(blob: bytes) -> ExtractedDoc:
    try:
        text_in = _decode(blob)
        rows = list(csv.reader(io.StringIO(text_in)))
        lines = ["\t".join(r) for r in rows if any(c.strip() for c in r)]
        text = "\n".join(lines).strip()
        return ExtractedDoc(
            text=text[:MAX_TEXT_CHARS],
            format="csv",
            mime_type=_MIME["csv"],
            ok=bool(text),
            pages=[text],
        )
    except Exception as exc:
        return ExtractedDoc(text="", format="csv", mime_type=_MIME["csv"], ok=False, error=str(exc))


def _extract_html(blob: bytes) -> ExtractedDoc:
    try:
        parser = _TextHTMLParser()
        parser.feed(_decode(blob))
        text = parser.text.strip()
        return ExtractedDoc(
            text=text[:MAX_TEXT_CHARS],
            format="html",
            mime_type=_MIME["html"],
            ok=bool(text),
            pages=[text],
        )
    except Exception as exc:
        return ExtractedDoc(
            text="", format="html", mime_type=_MIME["html"], ok=False, error=str(exc)
        )


def _extract_txt(blob: bytes) -> ExtractedDoc:
    text = _decode(blob).strip()
    return ExtractedDoc(
        text=text[:MAX_TEXT_CHARS],
        format="txt",
        mime_type=_MIME["txt"],
        ok=bool(text),
        pages=[text],
    )


def extract_text(filename: str, blob: bytes, *, allow_ocr: bool = True) -> ExtractedDoc:
    """Extract text from a single (non-archive) document. Archives must be
    expanded via ``documents.archive`` first — this returns ``unsupported`` for
    a ZIP so the caller doesn't accidentally treat it as text."""
    fmt = detect_format(filename, blob)
    if fmt == "pdf":
        return _extract_pdf(blob, allow_ocr=allow_ocr)
    if fmt == "docx":
        return _extract_docx(blob)
    if fmt == "xlsx":
        return _extract_xlsx(blob)
    if fmt == "csv":
        return _extract_csv(blob)
    if fmt == "html":
        return _extract_html(blob)
    if fmt == "txt":
        return _extract_txt(blob)
    return ExtractedDoc(
        text="",
        format=fmt,
        mime_type=_MIME.get(fmt, _MIME["unknown"]),
        ok=False,
        error=f"unsupported format: {fmt}",
    )
